"""
📚 Document Processor - The Smart Document Reader! 📖

This service is like a super-powered scanner that can:
1. 📄 Read ANY type of document (PDF, Word, HTML, images, etc.)
2. 🔍 Extract all the text content from them
3. ✂️ Break long documents into smaller, manageable pieces
4. 🏷️ Understand the structure and format of documents
5. 🔢 Calculate unique fingerprints (checksums) for each document

SUPPORTED FORMATS:
- PDF files (.pdf) - Most common document format
- Word documents (.docx) - Microsoft Word files
- HTML files (.html) - Web pages
- Markdown files (.md) - Simple text with formatting
- Plain text (.txt) - Simple text files
- Images (.jpg, .png) - Extract text using OCR (Optical Character Recognition)

HOW IT WORKS:
1. Document Detection: Figure out what type of file you have
2. Text Extraction: Get all the readable text from the document
3. Intelligent Chunking: Break text into logical pieces
4. Metadata Creation: Store information about the document
5. Quality Control: Ensure everything is processed correctly

This is the FIRST STEP in the RAG pipeline - without this, we can't read your documents!
"""
import os
import hashlib
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import logging

# Import all the tools we need to read different file types
import PyPDF2          # For reading PDF files
from bs4 import BeautifulSoup  # For reading HTML files
import markdown        # For reading Markdown files
from docx import Document as DocxDocument  # For reading Word documents
from PIL import Image  # For reading images
import pytesseract     # For extracting text from images (OCR)

# We used to use NLTK for text processing, but now we use simpler methods
# This makes the system faster and easier to set up

# Import our own tools and settings
from app.core.config import settings
from app.models.schemas import Document, Chunk, DocumentFormat, ChunkType

# Set up logging to track what's happening
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    🚀 Main Document Processor - The Multi-Format Reader!
    
    This class is like a universal translator for documents:
    - It can read ANY type of file you give it
    - It extracts all the text content
    - It breaks documents into logical pieces
    - It maintains the structure and meaning
    
    Think of it as a super-smart librarian who can read any language or format!
    """
    
    def __init__(self):
        """
        🏗️ Initialize the document processor - Set up our tools!
        
        We configure two important settings:
        - chunk_size: How big each piece of text should be (default: 500 words)
        - chunk_overlap: How much overlap between pieces (default: 100 words)
        
        Overlap is important because it helps maintain context between chunks!
        """
        # Get settings from our configuration
        self.chunk_size = settings.chunk_size      # How big each chunk should be
        self.chunk_overlap = settings.chunk_overlap  # How much overlap between chunks
        
        # We used to download NLTK data, but now we use simpler text splitting
        # This makes the system faster and easier to set up
    
    def process_document(self, file_path: str, title: str, metadata: Dict[str, Any] = None) -> Document:
        """
        📄 Process a document - The first step in understanding your file!
        
        This method is like examining a document before putting it in a filing cabinet:
        1. 📁 Check if the file actually exists
        2. 🔍 Figure out what type of file it is (PDF, Word, etc.)
        3. 🔢 Calculate a unique fingerprint (checksum) to identify it
        4. 🏷️ Create a record with all the basic information
        5. 🆔 Give it a unique ID for future reference
        
        Args:
            file_path: Where the document is located on your computer
            title: What to call this document
            metadata: Extra information about the document (optional)
        
        Returns:
            Document: A complete record of the document with all its information
        
        Think of this as creating a "library card" for your document!
        """
        # Convert the file path to a Path object for easier handling
        file_path = Path(file_path)
        
        print(f"📄 Processing document: {file_path}")
        print(f"✅ File exists: {file_path.exists()}")
        
        # STEP 1: Check if the file actually exists
        # This prevents errors when trying to process non-existent files
        if not file_path.exists():
            raise FileNotFoundError(f"❌ Document not found: {file_path}")
        
        # STEP 2: Figure out what type of file this is
        # Different file types need different processing methods
        print(f"🔍 Detecting document format...")
        doc_format = self._detect_format(file_path)
        print(f"📋 Detected format: {doc_format}")
        
        # STEP 3: Calculate a unique fingerprint (checksum)
        # This helps us know if the document has changed
        # Think of it like a DNA fingerprint for your document
        print(f"🔢 Calculating document fingerprint...")
        checksum = self._calculate_checksum(file_path)
        print(f"🆔 Calculated checksum: {checksum[:10]}...")
        
        # STEP 4: Create a complete document record
        # This is like creating a library card with all the important information
        print(f"📝 Creating document record...")
        document = Document(
            id=self._generate_id(),        # Generate a unique ID
            title=title,                   # The document title
            file_path=str(file_path),      # Where the file is located
            format=doc_format,             # What type of file it is
            checksum=checksum,             # The unique fingerprint
            metadata=metadata or {}        # Any extra information
        )
        
        print(f"✅ Created document with ID: {document.id}")
        print(f"🎉 Document '{title}' successfully processed!")
        return document
    
    def extract_text(self, document: Document) -> str:
        """Extract text content from document based on format."""
        file_path = Path(document.file_path)
        print(f"Extracting text from: {file_path}")
        print(f"Document format: {document.format}")
        
        try:
            if document.format == DocumentFormat.PDF:
                text = self._extract_pdf_text(file_path)
            elif document.format == DocumentFormat.HTML:
                text = self._extract_html_text(file_path)
            elif document.format == DocumentFormat.MARKDOWN:
                text = self._extract_markdown_text(file_path)
            elif document.format == DocumentFormat.DOCX:
                text = self._extract_docx_text(file_path)
            elif document.format == DocumentFormat.TEXT:
                text = self._extract_plain_text(file_path)
            elif document.format == DocumentFormat.IMAGE:
                text = self._extract_image_text(file_path)
            else:
                raise ValueError(f"Unsupported document format: {document.format}")
            
            print(f"Extracted text length: {len(text)}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {document.file_path}: {e}")
            print(f"Error extracting text: {e}")
            raise
    
    def chunk_text(self, text: str, document_id: str) -> List[Chunk]:
        """
        ✂️ Split text into intelligent, searchable chunks!
        
        This method uses a smart approach to break down long documents:
        1. 🔍 First, try to find natural breaks (headings, sections)
        2. 📏 If chunks are too long, break them down further
        3. 🔄 Use overlapping chunks to maintain context
        4. 🏷️ Label each chunk with metadata for better searching
        
        Args:
            text: The full text content to chunk
            document_id: Which document this text belongs to
        
        Returns:
            List of Chunk objects, each containing a piece of the text
        """
        print(f"✂️ Starting text chunking for document {document_id}")
        print(f"📊 Text length: {len(text)} characters, {len(text.split())} words")
        
        chunks = []
        
        # STEP 1: Try semantic chunking by headings and natural breaks
        print(f"🔍 Attempting semantic chunking...")
        semantic_chunks = self._semantic_chunking(text)
        print(f"✅ Semantic chunking found {len(semantic_chunks)} initial chunks")
        
        # STEP 2: Process each semantic chunk
        for i, (content, heading_path, start_pos) in enumerate(semantic_chunks):
            print(f"🔍 Processing chunk {i+1}/{len(semantic_chunks)}")
            
            # If chunk is too long, split it further using sliding window
            if len(content.split()) > self.chunk_size:
                print(f"📏 Chunk {i+1} is too long ({len(content.split())} words), splitting...")
                sub_chunks = self._sliding_window_chunking(content, start_pos)
                print(f"✂️ Split into {len(sub_chunks)} sub-chunks")
                chunks.extend(sub_chunks)
            else:
                # Chunk is the right size, create it directly
                print(f"✅ Chunk {i+1} is the right size ({len(content.split())} words)")
                chunk = self._create_chunk(
                    content, document_id, i, heading_path, start_pos, len(content)
                )
                chunks.append(chunk)
        
        print(f"🎉 Text chunking complete! Created {len(chunks)} total chunks")
        return chunks
    
    def _detect_format(self, file_path: Path) -> DocumentFormat:
        """Detect document format from file extension."""
        suffix = file_path.suffix.lower()
        
        format_mapping = {
            '.pdf': DocumentFormat.PDF,
            '.html': DocumentFormat.HTML,
            '.htm': DocumentFormat.HTML,
            '.md': DocumentFormat.MARKDOWN,
            '.markdown': DocumentFormat.MARKDOWN,
            '.txt': DocumentFormat.TEXT,
            '.docx': DocumentFormat.DOCX,
            '.png': DocumentFormat.IMAGE,
            '.jpg': DocumentFormat.IMAGE,
            '.jpeg': DocumentFormat.IMAGE,
            '.py': DocumentFormat.CODE,
            '.java': DocumentFormat.CODE,
            '.js': DocumentFormat.CODE,
            '.ts': DocumentFormat.CODE,
            '.cpp': DocumentFormat.CODE,
            '.c': DocumentFormat.CODE,
        }
        
        return format_mapping.get(suffix, DocumentFormat.TEXT)
    
    def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate SHA-256 checksum of file."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _generate_id(self) -> str:
        """Generate unique document ID."""
        import uuid
        return str(uuid.uuid4())
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.warning(f"PDF text extraction failed, trying OCR: {e}")
            if settings.enable_ocr:
                text = self._extract_pdf_ocr(file_path)
        
        return text.strip()
    
    def _extract_pdf_ocr(self, file_path: Path) -> str:
        """Extract text from PDF using OCR."""
        # This is a simplified OCR implementation
        # In production, you might want to use more sophisticated OCR
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            text = ""
            for page in doc:
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text += pytesseract.image_to_string(img) + "\n"
            return text
        except ImportError:
            logger.warning("PyMuPDF not available, OCR extraction failed")
            return ""
    
    def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text and clean it up
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            # For semantic chunking, we want to preserve heading structure
            # So we'll use the raw markdown content instead of converting to HTML
            return content
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    def _extract_plain_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_image_text(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        if not settings.enable_ocr:
            return ""
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {e}")
            return ""
    
    def _semantic_chunking(self, text: str) -> List[Tuple[str, List[str], int]]:
        """
        🔍 Split text by semantic boundaries (headings, sections, paragraphs).
        
        This method tries to find natural breaks in the text like:
        - Headings (## Title)
        - Paragraph breaks
        - Section dividers
        - Numbered lists
        
        If no clear boundaries are found, it falls back to simple paragraph splitting.
        """
        chunks = []
        
        # Split by common heading patterns
        heading_patterns = [
            r'^#{1,6}\s+(.+)$',  # Markdown headings (# ## ### etc.)
            r'^[A-Z][A-Z\s]+\n[-=]+\n',  # Underlined headings
            r'^\d+\.\s+[A-Z]',  # Numbered sections (1. Title)
            r'^[A-Z][a-z]+\s+\d+',  # Chapter titles (Chapter 1)
        ]
        
        lines = text.split('\n')
        current_chunk = []
        current_heading_path = []
        start_pos = 0
        
        for i, line in enumerate(lines):
            is_heading = any(re.match(pattern, line) for pattern in heading_patterns)
            
            if is_heading and current_chunk:
                # Save current chunk
                chunk_text = '\n'.join(current_chunk).strip()
                if chunk_text:
                    chunks.append((chunk_text, current_heading_path.copy(), start_pos))
                
                # Start new chunk
                current_chunk = [line]
                current_heading_path = [line.strip()]
                start_pos = i
            else:
                current_chunk.append(line)
        
        # Add final chunk
        if current_chunk:
            chunk_text = '\n'.join(current_chunk).strip()
            if chunk_text:
                chunks.append((chunk_text, current_heading_path, start_pos))
        
        # If no chunks were created (no headings found), fall back to paragraph-based chunking
        if not chunks:
            print(f"🔍 No headings found, using paragraph-based chunking...")
            chunks = self._paragraph_based_chunking(text)
        
        print(f"🔍 Semantic chunking created {len(chunks)} chunks")
        return chunks
    
    def _paragraph_based_chunking(self, text: str) -> List[Tuple[str, List[str], int]]:
        """
        📝 Split text into chunks based on paragraphs and sentence boundaries.
        
        This is a fallback method when no clear headings are found.
        It creates chunks by:
        1. Splitting on double newlines (paragraphs)
        2. If paragraphs are too long, splitting on sentences
        3. If sentences are too long, using sliding window
        """
        chunks = []
        
        # Split by paragraphs (double newlines)
        paragraphs = text.split('\n\n')
        
        for i, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # If paragraph is short enough, make it a chunk
            if len(paragraph.split()) <= self.chunk_size:
                chunks.append((paragraph, [], i))
            else:
                # If paragraph is too long, split by sentences
                sentences = self._split_into_sentences(paragraph)
                for j, sentence in enumerate(sentences):
                    if sentence.strip():
                        chunks.append((sentence.strip(), [], i * 1000 + j))
        
        print(f"📝 Paragraph-based chunking created {len(chunks)} chunks")
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using simple punctuation rules."""
        # Simple sentence splitting - split on common sentence endings
        import re
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _sliding_window_chunking(self, text: str, start_offset: int) -> List[Chunk]:
        """
        🔄 Split text using sliding window approach.
        
        This method creates overlapping chunks to maintain context:
        - Each chunk has a maximum size (chunk_size)
        - Chunks overlap by a certain amount (chunk_overlap)
        - This ensures no information is lost between chunks
        """
        chunks = []
        words = text.split()  # Simple word splitting instead of NLTK
        
        if len(words) <= self.chunk_size:
            return []
        
        for i in range(0, len(words) - self.chunk_size + 1, self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunk = self._create_chunk(
                chunk_text, 
                "",  # document_id will be set later
                len(chunks), 
                [], 
                start_offset + i, 
                len(chunk_text)
            )
            chunks.append(chunk)
        
        return chunks
    
    def _create_chunk(self, content: str, document_id: str, chunk_index: int, 
                      heading_path: List[str], start_offset: int, end_offset: int) -> Chunk:
        """
        🧩 Create a chunk object with all necessary information.
        
        A chunk is a piece of text that can be:
        - Searched independently
        - Embedded into a vector
        - Retrieved when relevant to a query
        
        Args:
            content: The actual text content
            document_id: Which document this chunk belongs to
            chunk_index: Position of this chunk in the document
            heading_path: Any headings that describe this chunk
            start_offset: Where this chunk starts in the original text
            end_offset: Where this chunk ends in the original text
        """
        # Generate a unique ID for this chunk
        chunk_id = f"{document_id}_chunk_{chunk_index}" if document_id else f"temp_chunk_{chunk_index}"
        checksum = hashlib.md5(content.encode()).hexdigest()
        
        # Determine what type of content this chunk contains
        chunk_type = ChunkType.TEXT  # Default to regular text
        if any(keyword in content.lower() for keyword in ['def ', 'class ', 'function ', 'import ']):
            chunk_type = ChunkType.CODE  # Contains programming code
        elif heading_path:
            chunk_type = ChunkType.HEADING  # Contains a heading
        
        # Create and return the chunk object
        return Chunk(
            id=chunk_id,
            document_id=document_id,
            content=content,
            chunk_type=chunk_type,
            heading_path=heading_path,
            start_offset=start_offset,
            end_offset=end_offset,
            token_count=len(content.split()),  # Count words as tokens
            checksum=checksum
        )
